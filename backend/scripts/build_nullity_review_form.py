from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

ROOT = Path(__file__).resolve().parents[2]
DATA_PATH = ROOT / "backend/data/nullity_benchmark_recognition.json"
OUTPUT_PATH = ROOT / "docs/ficha-revisao-guiada-reconhecimento.docx"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "203040"
MUTED = "5F6B76"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
CAUTION = "FFF4CE"
WHITE = "FFFFFF"
LINE = "C8D0D8"

REQUIREMENTS = [
    "A pessoa era desconhecida antes do fato",
    "Houve descrição antes da exibição",
    "Foram mostradas pessoas ou fotos semelhantes",
    "A polícia evitou sugerir uma resposta",
    "O ato foi registrado em termo, ata ou gravação",
    "As testemunhas foram ouvidas separadamente",
    "Não houve repetição indevida do reconhecimento",
    "Existe outra prova de autoria",
]

TEXT_REPLACEMENTS = {
    "A vitima": "A vítima",
    "a vitima": "a vítima",
    "A primeira vitima": "A primeira vítima",
    "A segunda vitima": "A segunda vítima",
    "As vitimas": "As vítimas",
    "das vitimas": "das vítimas",
    "depoimento vitima": "depoimento da vítima",
    "auto reconhecimento": "auto de reconhecimento",
    "vitimas": "vítimas",
    "vitima": "vítima",
    "nao": "não",
    "Nao": "Não",
    "acusacao": "acusação",
    "atribuicao": "atribuição",
    "descricao": "descrição",
    "caracteristicas": "características",
    "identificacao": "identificação",
    "Identificacao": "Identificação",
    "investigacao": "investigação",
    "pericia": "perícia",
    "impressao": "impressão",
    "gravacao": "gravação",
    "camera": "câmera",
    "unica": "única",
    "previa": "prévia",
    "deposito": "depósito",
    "custodia": "custódia",
    "autonoma": "autônoma",
    "denuncia": "denúncia",
    "ja": "já",
    "confissao": "confissão",
    "delegacia": "delegacia",
    "apresentacao": "apresentação",
    "comunicacao": "comunicação",
    "ameaca": "ameaça",
    "intimidacao": "intimidação",
    "biologico": "biológico",
    "propria": "própria",
    "derivacao": "derivação",
    "investigado": "investigado",
}


def main() -> None:
    payload = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    document = Document()
    _configure_document(document)
    _add_intro_page(document, payload)
    for index, case in enumerate(payload["cases"], start=1):
        _add_case_page(document, case, index, len(payload["cases"]))
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.2677)
    section.page_height = Inches(11.6929)
    section.top_margin = Inches(0.75)
    section.right_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Title", 23, INK, 0, 4),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    document.core_properties.title = "Ficha guiada de revisão do reconhecimento de pessoas"
    document.core_properties.subject = "Validação jurídica do benchmark de nulidades"
    document.core_properties.author = "Projeto Preparador de Audiência"
    document.core_properties.keywords = "revisão jurídica, reconhecimento, nulidade"

    _configure_footer(section, total_pages=7)


def _configure_footer(section, total_pages: int) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Página ")
    _set_run(run, size=8.5, color=MUTED)
    _add_page_field(paragraph)
    run = paragraph.add_run(f" de {total_pages}")
    _set_run(run, size=8.5, color=MUTED)


def _add_intro_page(document: Document, payload: dict[str, object]) -> None:
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(6)
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run("FICHA DE REVISÃO GUIADA")
    _set_run(run, size=10, color=BLUE, bold=True)

    title = document.add_paragraph(style="Title")
    title.add_run("O reconhecimento de pessoas foi feito corretamente?")

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run(
        "Seis casos curtos para comparar a análise da ferramenta com a avaliação de um profissional"
    )
    _set_run(run, size=12, color=MUTED)

    _add_callout(
        document,
        "O que você precisa fazer",
        (
            "Leia os trechos de cada caso e marque as respostas que melhor descrevem o que "
            "aconteceu. Não é necessário escrever um parecer, calcular pena ou citar artigos."
        ),
        fill=CAUTION,
    )

    document.add_heading("Quem está revisando", level=2)
    _add_labeled_line(document, "Nome")
    _add_labeled_line(document, "Cargo ou função")
    _add_labeled_line(document, "Data da revisão")

    document.add_heading("Como funciona", level=2)
    steps = [
        "Leia somente os trechos apresentados no início da página.",
        "Marque uma resposta em cada pergunta. Use “não informado” quando o trecho não disser.",
        "Anote a página que sustenta sua conclusão.",
        "No final, explique em uma frase o motivo da sua resposta ou indique o que faltou.",
    ]
    for index, text in enumerate(steps, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.first_line_indent = Inches(-0.18)
        paragraph.paragraph_format.space_after = Pt(4)
        number = paragraph.add_run(f"{index}. ")
        _set_run(number, bold=True, color=DARK_BLUE)
        paragraph.add_run(text)

    document.add_heading("Três expressões usadas na ficha", level=2)
    definitions = [
        ("Reconhecimento", "quando vítima ou testemunha aponta quem acredita ser o autor."),
        ("Outra prova", "prova que liga o acusado ao fato sem depender do reconhecimento."),
        ("Inválido ou nulo", "o ato pode não ser aproveitado porque o procedimento teve falhas."),
        ("Não informado", "os trechos apresentados não permitem responder."),
    ]
    table = document.add_table(rows=0, cols=2)
    _set_table_geometry(table, [2700, 6660])
    for label, meaning in definitions:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = meaning
        _format_definition_row(cells)
    _set_table_geometry(table, [2700, 6660])

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run(
        "As respostas do sistema não aparecem nesta ficha para não influenciar a revisão."
    )
    _set_run(run, size=9, color=MUTED, italic=True)


def _add_case_page(
    document: Document,
    case: dict[str, object],
    index: int,
    total: int,
) -> None:
    kicker = document.add_paragraph()
    kicker.paragraph_format.page_break_before = True
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run(f"Caso {index} de {total}")
    _set_run(run, size=9, color=BLUE, bold=True)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title.paragraph_format.keep_with_next = True
    run = title.add_run(_humanize(str(case["title"])))
    _set_run(run, size=15, color=INK, bold=True)

    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(4)
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run("Trechos apresentados")
    _set_run(run, size=11.5, color=DARK_BLUE, bold=True)

    sources = case["sources"]
    compact_sources = len(sources) > 2
    for source in sources:
        _add_source_block(document, source, compact=compact_sources)

    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(5)
    heading.paragraph_format.space_after = Pt(3)
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run("Sua análise")
    _set_run(run, size=11.5, color=DARK_BLUE, bold=True)

    _add_option_line(
        document,
        "1. O caso contém um reconhecimento de pessoa",
        ["Sim", "Não", "Só menciona, sem explicar", "Não sei"],
        wrap_after=2,
    )
    _add_option_line(
        document,
        "2. A pessoa já era conhecida pela vítima ou testemunha",
        ["Sim", "Não", "Não informado"],
    )
    _add_option_line(
        document,
        "3. Como o reconhecimento foi feito",
        [
            "Várias pessoas ou fotos",
            "Uma foto isolada",
            "O trecho não explica",
            "Não houve reconhecimento",
        ],
        wrap_after=2,
    )
    _add_option_line(
        document,
        "4. Existe outra prova que liga o acusado ao fato",
        ["Sim", "Não", "Não informado"],
    )
    _add_option_line(
        document,
        "5. Qual é a sua conclusão",
        ["Parece regular", "Pode ser inválido ou nulo", "Faltam informações", "Não se aplica"],
        wrap_after=2,
    )
    _add_option_line(
        document,
        "6. Sem esse reconhecimento, ainda há prova contra o acusado",
        ["Sim", "Não", "Não dá para saber"],
    )

    caption = document.add_paragraph()
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(3)
    caption.paragraph_format.keep_with_next = True
    run = caption.add_run("Conferência rápida dos fatos")
    _set_run(run, size=10, color=DARK_BLUE, bold=True)

    explanation = document.add_paragraph()
    explanation.paragraph_format.space_after = Pt(3)
    run = explanation.add_run(
        "Marque Sim quando o trecho confirmar a frase, Não quando ele a contradizer e "
        "Não informado quando não houver dados suficientes."
    )
    _set_run(run, size=8.2, color=MUTED, italic=True)

    _add_requirements_table(document)

    _add_option_line(
        document,
        "7. Quão seguro você está da conclusão",
        ["Alta", "Média", "Baixa"],
    )

    comments = document.add_paragraph()
    comments.paragraph_format.space_before = Pt(3)
    comments.paragraph_format.space_after = Pt(2)
    comments.paragraph_format.keep_with_next = True
    run = comments.add_run("Explique em uma frase o motivo ou diga o que faltou")
    _set_run(run, size=9.5, color=DARK_BLUE, bold=True)
    _add_response_box(document)


def _add_source_block(
    document: Document,
    source: dict[str, object],
    *,
    compact: bool = False,
) -> None:
    table = document.add_table(rows=1, cols=1)
    _set_table_geometry(table, [PAGE_WIDTH_DXA])
    cell = table.cell(0, 0)
    _set_cell_shading(cell, LIGHT_GRAY)
    _set_cell_border(cell, LINE)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    label = paragraph.add_run(
        f"Página {source['page']}  |  {_document_type(source.get('document_type'))}"
    )
    _set_run(label, size=7.8 if compact else 8.2, color=BLUE, bold=True)
    text_paragraph = cell.add_paragraph()
    text_paragraph.paragraph_format.space_before = Pt(0)
    text_paragraph.paragraph_format.space_after = Pt(0)
    text_paragraph.paragraph_format.line_spacing = 0.95 if compact else 1.05
    run = text_paragraph.add_run(_humanize(str(source["text"])))
    _set_run(run, size=8.1 if compact else 8.7, color=INK)
    if not compact:
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(2)


def _add_requirements_table(document: Document) -> None:
    widths = [5000, 900, 900, 1400, 1160]
    table = document.add_table(rows=1, cols=5)
    _set_table_geometry(table, widths)
    headers = ["O que os trechos mostram", "Sim", "Não", "Não informado", "Página"]
    for cell, text in zip(table.rows[0].cells, headers, strict=True):
        cell.text = text
        _set_cell_shading(cell, LIGHT_BLUE)
        _set_cell_border(cell, LINE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                _set_run(run, size=7.2, color=DARK_BLUE, bold=True)
    _repeat_table_header(table.rows[0])

    for requirement in REQUIREMENTS:
        cells = table.add_row().cells
        values = [requirement, "☐", "☐", "☐", ""]
        for index, (cell, value) in enumerate(zip(cells, values, strict=True)):
            cell.text = value
            _set_cell_border(cell, LINE)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if index == 0
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    _set_run(run, size=7.8 if index == 0 else 8, color=INK)
    _set_table_geometry(table, widths)


def _add_option_line(
    document: Document,
    label: str,
    options: list[str],
    *,
    wrap_after: int | None = None,
) -> None:
    groups = [options]
    if wrap_after is not None:
        groups = [options[:wrap_after], options[wrap_after:]]
    for group_index, group in enumerate(groups):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
        if group_index == 0:
            label_run = paragraph.add_run(f"{label}: ")
            _set_run(label_run, size=8.8, color=INK, bold=True)
        else:
            paragraph.paragraph_format.left_indent = Inches(0.2)
        for option in group:
            option_run = paragraph.add_run(f"☐ {option}    ")
            _set_run(option_run, size=8.5, color=INK)


def _add_response_box(document: Document) -> None:
    table = document.add_table(rows=1, cols=1)
    _set_table_geometry(table, [PAGE_WIDTH_DXA])
    cell = table.cell(0, 0)
    _set_cell_shading(cell, WHITE)
    _set_cell_border(cell, LINE)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.add_run(" ")
    blank = cell.add_paragraph(" ")
    blank.paragraph_format.space_before = Pt(0)
    blank.paragraph_format.space_after = Pt(5)


def _add_callout(
    document: Document,
    label: str,
    text: str,
    *,
    fill: str,
) -> None:
    table = document.add_table(rows=1, cols=1)
    _set_table_geometry(table, [PAGE_WIDTH_DXA])
    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill)
    _set_cell_border(cell, LINE)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    label_run = paragraph.add_run(f"{label}: ")
    _set_run(label_run, size=10, color=INK, bold=True)
    text_run = paragraph.add_run(text)
    _set_run(text_run, size=10, color=INK)


def _add_labeled_line(document: Document, label: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(f"{label}: ")
    _set_run(run, size=10.5, color=INK, bold=True)
    line = paragraph.add_run("____________________________________________________")
    _set_run(line, size=10.5, color=MUTED)


def _format_definition_row(cells) -> None:
    for index, cell in enumerate(cells):
        _set_cell_border(cell, LINE)
        _set_cell_shading(cell, LIGHT_BLUE if index == 0 else WHITE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                _set_run(
                    run,
                    size=9.5,
                    color=DARK_BLUE if index == 0 else INK,
                    bold=index == 0,
                )


def _set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.allow_autofit = False
    if table.rows:
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table_element = table._tbl
    properties = table_element.tblPr

    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths_dxa)))
    width.set(qn("w:type"), "dxa")

    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")

    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    _set_table_cell_margins(properties)

    grid = table_element.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_dxa in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width_dxa))
        grid.append(column)

    for row in table.rows:
        for cell, width_dxa in zip(row.cells, widths_dxa, strict=True):
            cell.width = Twips(width_dxa)
            cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:w"), str(width_dxa))
            cell_width.set(qn("w:type"), "dxa")


def _set_table_cell_margins(properties) -> None:
    margins = properties.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        properties.append(margins)
    for side, value in [("top", 70), ("bottom", 70), ("start", 120), ("end", 120)]:
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_cell_border(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def _repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _set_run(
    run,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    _set_run(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])


def _document_type(value: object) -> str:
    if not value:
        return "Trecho processual"
    return _humanize(str(value).replace("_", " ")).capitalize()


def _humanize(text: str) -> str:
    result = text
    for source, replacement in TEXT_REPLACEMENTS.items():
        result = re.sub(rf"\b{re.escape(source)}\b", replacement, result)
    return result


if __name__ == "__main__":
    main()
